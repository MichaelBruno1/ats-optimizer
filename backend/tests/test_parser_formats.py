"""Unit tests for robust document extraction across various file formats (PDF, DOCX, TXT with BOM)."""

import io
import pytest
import fitz
from docx import Document

from app.services.document_parser import (
    extract_text_from_bytes,
    DocumentParseError,
)


def test_pdf_extraction_success() -> None:
    """Should correctly extract text from a programmatically generated PDF."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Resume Candidate: Jane Doe")
    page.insert_text((72, 100), "Skills: Python, FastAPI, Docker")
    pdf_bytes = doc.write()
    doc.close()

    extracted = extract_text_from_bytes(pdf_bytes, "resume.pdf")
    assert "Jane Doe" in extracted
    assert "FastAPI" in extracted


def test_docx_extraction_paragraphs_and_tables() -> None:
    """Should extract text from DOCX paragraphs, tables, headers, and footers."""
    doc = Document()
    
    # Add a section header/footer
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Header Info: Confidential"
    footer = section.footer
    footer.paragraphs[0].text = "Footer Info: Page 1"

    # Add standard body paragraphs
    doc.add_paragraph("Work Experience:")
    doc.add_paragraph("Senior Python Developer at TechCorp")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Java"
    table.cell(0, 1).text = "AWS"
    table.cell(1, 0).text = "2020-2024"
    table.cell(1, 1).text = "Optimized backend pipelines"

    bio = io.BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()

    extracted = extract_text_from_bytes(docx_bytes, "resume.docx")
    
    # Assert header and footer content is extracted
    assert "Header Info: Confidential" in extracted
    assert "Footer Info: Page 1" in extracted
    
    # Assert main paragraphs are extracted
    assert "Work Experience:" in extracted
    assert "Senior Python Developer" in extracted
    
    # Assert table contents are extracted
    assert "Java | AWS" in extracted
    assert "2020-2024 | Optimized backend pipelines" in extracted


def test_txt_utf8_bom_handling() -> None:
    """Should automatically strip UTF-8 BOM when decoding plain text files."""
    # UTF-8 BOM is \xef\xbb\xbf
    bom_content = b"\xef\xbb\xbfName: Bob Smith\nRole: Data Engineer"
    extracted = extract_text_from_bytes(bom_content, "resume.txt")
    
    # Make sure BOM is stripped and name is resolved
    assert extracted.startswith("Name:")
    assert "Bob Smith" in extracted
    assert "Data Engineer" in extracted


def test_null_byte_stripping() -> None:
    """Should strip null bytes from the extracted text across all formats."""
    # A TXT file containing null bytes
    null_byte_content = "John\x00 Doe\x00\nSoftware\x00 Engineer".encode("utf-8")
    extracted = extract_text_from_bytes(null_byte_content, "resume.txt")
    
    assert "\x00" not in extracted
    assert "John Doe" in extracted
    assert "Software Engineer" in extracted
