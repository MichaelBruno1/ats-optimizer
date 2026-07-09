"""Document parser service.

Handles extraction of plain text from uploaded resume files.
Supported formats: PDF (.pdf via PyMuPDF), DOCX (.docx via python-docx),
and plain text (.txt).

All validation (size, format) is centralised here so the API layer remains thin.
"""

import logging
from pathlib import Path

from fastapi import UploadFile

from app.config import settings

logger = logging.getLogger(__name__)

ACCEPTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".txt"}
ACCEPTED_MIME_TYPES: set[str] = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed due to format or content issues."""


class FileTooLargeError(ValueError):
    """Raised when an uploaded file exceeds the configured size limit."""


class UnsupportedFormatError(ValueError):
    """Raised when the uploaded file has an unsupported extension or MIME type."""


# ─────────────────────────────────────────────────────────────────────────────
# Public interface
# ─────────────────────────────────────────────────────────────────────────────


async def extract_text_from_upload(upload: UploadFile) -> str:
    """Validate and extract plain text from a FastAPI UploadFile.

    Performs three validation steps before extraction:
    1. File size ≤ MAX_FILE_SIZE_MB.
    2. File extension in ACCEPTED_EXTENSIONS.
    3. Successful text extraction.

    Args:
        upload: The FastAPI UploadFile received from the multipart request.

    Returns:
        Extracted plain text from the document.

    Raises:
        FileTooLargeError: If the file exceeds the configured size limit.
        UnsupportedFormatError: If the file extension is not supported.
        DocumentParseError: If text extraction fails for any reason.
    """
    raw_bytes = await upload.read()

    # ── Size validation ───────────────────────────────────────────────────────
    if len(raw_bytes) > settings.max_file_size_bytes:
        raise FileTooLargeError(
            f"File size {len(raw_bytes) / 1_048_576:.1f} MB exceeds the "
            f"{settings.max_file_size_mb} MB limit."
        )

    # ── Extension validation ──────────────────────────────────────────────────
    filename = upload.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ACCEPTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported file format '{ext}'. "
            f"Accepted formats: {', '.join(sorted(ACCEPTED_EXTENSIONS))}."
        )

    # ── Text extraction (Offloaded to executor to avoid event loop blocking) ──
    logger.debug("Extracting text from '%s' (%d bytes) in executor", filename, len(raw_bytes))
    import asyncio
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(
        None, _dispatch_extraction, raw_bytes, ext, filename
    )


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extract text from raw bytes given a filename (for testing purposes).

    Args:
        content: Raw file bytes.
        filename: Original filename used to determine the format.

    Returns:
        Extracted plain text.

    Raises:
        UnsupportedFormatError: If the extension is not supported.
        DocumentParseError: If extraction fails.
    """
    ext = Path(filename).suffix.lower()
    if ext not in ACCEPTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported file format '{ext}'. "
            f"Accepted formats: {', '.join(sorted(ACCEPTED_EXTENSIONS))}."
        )
    return _dispatch_extraction(content, ext, filename)


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────


def _dispatch_extraction(raw_bytes: bytes, ext: str, filename: str) -> str:
    """Route extraction to the appropriate handler based on file extension.

    Args:
        raw_bytes: Raw file content.
        ext: Lowercase file extension (e.g. '.pdf').
        filename: Original filename (used only in error messages).

    Returns:
        Extracted plain text string.

    Raises:
        DocumentParseError: If the extraction fails.
    """
    try:
        if ext == ".pdf":
            text = _extract_from_pdf(raw_bytes)
        elif ext == ".docx":
            text = _extract_from_docx(raw_bytes)
        elif ext == ".txt":
            text = _extract_from_txt(raw_bytes)
        else:
            # Unreachable after format validation, but makes the type checker happy.
            raise UnsupportedFormatError(f"No extractor for extension '{ext}'.")
        return text.replace("\x00", "").strip()
    except (UnsupportedFormatError, FileTooLargeError):
        raise
    except Exception as exc:
        logger.exception("Failed to extract text from '%s'", filename)
        raise DocumentParseError(
            f"Could not extract text from '{filename}': {exc}"
        ) from exc


def _extract_from_pdf(raw_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF.

    Args:
        raw_bytes: Raw PDF file bytes.

    Returns:
        Concatenated plain text from all pages.

    Raises:
        DocumentParseError: If PyMuPDF cannot open or parse the document.
    """
    import fitz  # PyMuPDF — imported lazily to keep startup fast

    try:
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
    except Exception as exc:
        raise DocumentParseError(f"PyMuPDF could not open the PDF: {exc}") from exc

    pages: list[str] = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pages.append(page.get_text("text"))  # type: ignore[arg-type]

    doc.close()
    text = "\n".join(pages).strip()

    if not text:
        raise DocumentParseError(
            "The PDF appears to be image-based or empty. "
            "No selectable text could be extracted."
        )

    logger.debug("Extracted %d characters from PDF (%d pages)", len(text), len(pages))
    return text


def _extract_from_docx(raw_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx.

    Includes paragraphs, tables, headers, and footers.

    Args:
        raw_bytes: Raw DOCX file bytes.

    Returns:
        Concatenated plain text.

    Raises:
        DocumentParseError: If python-docx cannot parse the document.
    """
    import io

    from docx import Document  # python-docx

    try:
        doc = Document(io.BytesIO(raw_bytes))
    except Exception as exc:
        raise DocumentParseError(f"python-docx could not open the file: {exc}") from exc

    text_parts: list[str] = []

    def extract_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

    def extract_tables(tables) -> None:
        for table in tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

    # Extract headers
    for section in doc.sections:
        if section.header:
            extract_paragraphs(section.header.paragraphs)
            extract_tables(section.header.tables)

    # Extract body paragraphs & tables
    extract_paragraphs(doc.paragraphs)
    extract_tables(doc.tables)

    # Extract footers
    for section in doc.sections:
        if section.footer:
            extract_paragraphs(section.footer.paragraphs)
            extract_tables(section.footer.tables)

    text = "\n".join(text_parts).strip()

    if not text:
        raise DocumentParseError(
            "The DOCX document appears to be empty. No text content found."
        )

    logger.debug("Extracted %d characters from DOCX", len(text))
    return text


def _extract_from_txt(raw_bytes: bytes) -> str:
    """Decode a plain text file, attempting UTF-8-sig (with BOM support) then latin-1 fallback.

    Args:
        raw_bytes: Raw text file bytes.

    Returns:
        Decoded string content.

    Raises:
        DocumentParseError: If decoding fails with both encodings.
    """
    for encoding in ("utf-8-sig", "latin-1"):
        try:
            text = raw_bytes.decode(encoding).strip()
            if not text:
                raise DocumentParseError("The text file is empty.")
            logger.debug(
                "Decoded TXT (%s, %d chars)", encoding, len(text)
            )
            return text
        except UnicodeDecodeError:
            continue

    raise DocumentParseError(
        "Could not decode the text file with UTF-8 or Latin-1 encoding."
    )
